import os
from fpdf import FPDF
from docx import Document
import uuid

EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

class ExportService:
    def generate_pdf(self, text: str) -> str:
        """Generates a PDF file from the given text and returns the file path."""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Title
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Summary Report", ln=True, align="C")
            pdf.ln(10)
            
            # Content
            pdf.set_font("Arial", size=12)
            # FPDF handling of unicode can be tricky, relying on latin-1 by default. 
            # We'll encode/decode to latin-1 to avoid errors, replacing unsupported chars.
            # A more robust solution involves loading detailed fonts, but for this quick impl:
            clean_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, clean_text)
            
            filename = f"summary_{uuid.uuid4().hex}.pdf"
            filepath = os.path.join(EXPORT_DIR, filename)
            pdf.output(filepath)
            return filepath
        except Exception as e:
            print(f"[EXPORT-ERROR] PDF Generation failed: {e}")
            raise e

    def generate_docx(self, text: str) -> str:
        """Generates a DOCX file from the given text and returns the file path."""
        try:
            doc = Document()
            doc.add_heading('Summary Report', 0)
            doc.add_paragraph(text)
            
            filename = f"summary_{uuid.uuid4().hex}.docx"
            filepath = os.path.join(EXPORT_DIR, filename)
            doc.save(filepath)
            return filepath
        except Exception as e:
            print(f"[EXPORT-ERROR] DOCX Generation failed: {e}")
            raise e

export_service = ExportService()
